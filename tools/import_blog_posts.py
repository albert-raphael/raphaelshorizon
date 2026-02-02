#!/usr/bin/env python3
"""

Import content from DOCX files into blog post HTML files.

- Reads content from .docx files in a specified source directory.
- Uses blog-template.html to generate the final HTML.
- Extracts the first paragraph as the title and the rest as the body.
- Assumes a corresponding image exists in `frontend/assets/images/` named `blog-post-N.png`.

Usage:
    python tools/import_blog_posts.py

Prerequisites:
    pip install python-docx

"""
import os
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: 'python-docx' is not installed. Please run 'pip install python-docx'")
    exit(1)

ROOT = Path(__file__).resolve().parent.parent
FRONTEND = ROOT / 'frontend'
BLOG_TEMPLATE_PATH = FRONTEND / 'pages' / 'blog' / 'blog-template.html'

# --- Configuration ---
# Directory to search for .docx articles
DOCS_SEARCH_DIR = ROOT / 'documents'
# The range of blog posts to process
POST_RANGE = range(1, 7)  # This will process post-1.docx through post-6.docx
# ---------------------

def find_docx_path(post_num):
    """
    Searches for a docx file for a given post number within DOCS_SEARCH_DIR.
    Returns the full path if found, otherwise None.
    """
    file_name_to_find = f'post-{post_num}.docx'
    # Walk through the documents directory
    for dirpath, _, filenames in os.walk(DOCS_SEARCH_DIR):
        for filename in filenames:
            if filename.lower() == file_name_to_find.lower():
                # Found it!
                return Path(dirpath) / filename
    return None

def generate_html_from_docx(doc_path):
    """
    Extracts title and content from a DOCX file and formats it as HTML.
    Handles bold, italic, and underline formatting.
    """
    def escape_html(text):
        """Basic HTML escaping for text content."""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def runs_to_html(runs):
        """Converts a list of docx.text.run.Run objects to an HTML string."""
        html_parts = []
        for run in runs:
            text = escape_html(run.text)
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            html_parts.append(text)
        return "".join(html_parts)

    if not doc_path.exists():
        return None, None, None

    doc = Document(doc_path)
    
    if not doc.paragraphs:
        return None, None, None

    # Assume the first non-empty paragraph is the title
    title = ""
    first_content_index = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            title = p.text.strip()
            first_content_index = i + 1
            break

    # The rest of the paragraphs form the content
    html_content = []
    for p in doc.paragraphs[first_content_index:]:
        if p.runs:  # Process only if there is content
            paragraph_html = runs_to_html(p.runs)
            html_content.append(f"<p>{paragraph_html}</p>")
    
    return title, "\n".join(html_content), "An article from my collection."


def main():
    """Main script execution."""
    try:
        with open(BLOG_TEMPLATE_PATH, 'r', encoding='utf-8') as fh:
            blog_template = fh.read()
    except FileNotFoundError:
        print(f"Error: Blog template not found at {BLOG_TEMPLATE_PATH}")
        return

    print("Starting blog post import...")

    for n in POST_RANGE:
        post_num = n
        
        # Find the docx file automatically
        doc_path = find_docx_path(post_num)
        if not doc_path:
            print(f"WARN: Could not find 'post-{post_num}.docx' anywhere inside '{DOCS_SEARCH_DIR.relative_to(ROOT)}'.")
            continue

        html_path = FRONTEND / 'pages' / 'blog' / f'post-{post_num}.html'

        title, body_html, subtitle = generate_html_from_docx(doc_path)

        if title and body_html:
            content = blog_template.replace('POST_TITLE', title)
            # Also update the <title> tag
            content = re.sub(r'<title>.*?</title>', f'<title>{title} — Raphael\'s Horizon</title>', content)
            content = content.replace('POST_SUBTITLE', subtitle)
            # Use a regex for more robust replacement of the placeholder content
            content = re.sub(
                r'<p>Start your blog post content here.*?tags.</p>', 
                body_html, 
                content, 
                flags=re.DOTALL
            )
            # Also update the main image, assuming a naming convention like `blog-post-1.png`
            content = re.sub(r'assets/images/blog-post-\d+\.(jpg|png|jpeg|gif)', f'assets/images/blog-post-{post_num}.png', content)

            with open(html_path, 'w', encoding='utf-8') as fh:
                fh.write(content)
            print(f"Successfully populated '{html_path.relative_to(ROOT)}' from '{doc_path.relative_to(ROOT)}'")
        else:
            print(f"WARN: Could not process '{doc_path.relative_to(ROOT)}'. File might be empty, not found, or has no content.")

    print("\nDone.")

if __name__ == '__main__':
    main()
